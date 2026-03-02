"""
File parser for document processing.

Supports parsing Word (.docx), PDF (.pdf), Markdown (.md), and plain text (.txt) files.

Author: RAG Team
Created: 2026-03-02
"""

from pathlib import Path
from typing import Dict, List, Optional

from docx import Document as WordDocument
from pymupdf import Document as PDFDocument

from utils.logger_config import get_logger

logger = get_logger(__name__)


class FileParseError(Exception):
    """Base exception for file parsing errors."""

    pass


class UnsupportedFileTypeError(FileParseError):
    """Raised when file type is not supported."""

    pass


class FileNotFoundError(FileParseError):
    """Raised when file is not found."""

    pass


class FileParser:
    """
    Parser for various document file formats.

    Supports:
    - Word (.docx)
    - PDF (.pdf)
    - Markdown (.md)
    - Plain text (.txt)
    """

    # Supported file extensions
    SUPPORTED_EXTENSIONS: Dict[str, str] = {
        ".docx": "word",
        ".pdf": "pdf",
        ".md": "markdown",
        ".txt": "text",
    }

    def __init__(self, allowed_extensions: Optional[List[str]] = None):
        """
        Initialize the file parser.

        Args:
            allowed_extensions: List of allowed file extensions (e.g., [".pdf", ".docx"])
        """
        self.allowed_extensions = allowed_extensions or list(self.SUPPORTED_EXTENSIONS.keys())

    def parse_file(self, file_path: str) -> str:
        """
        Parse a file and return its text content.

        Args:
            file_path: Path to the file to parse

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            UnsupportedFileTypeError: If file type is not supported
            FileParseError: If parsing fails
        """
        # Check if file exists
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get file extension
        ext = path.suffix.lower()

        # Check if extension is supported
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {ext}. Supported types: {list(self.SUPPORTED_EXTENSIONS.keys())}"
            )

        # Check if extension is allowed
        if ext not in self.allowed_extensions:
            raise UnsupportedFileTypeError(
                f"File type {ext} is not in the allowed list: {self.allowed_extensions}"
            )

        # Parse based on file type
        file_type = self.SUPPORTED_EXTENSIONS[ext]

        try:
            if file_type == "word":
                return self._parse_word(file_path)
            elif file_type == "pdf":
                return self._parse_pdf(file_path)
            elif file_type == "markdown":
                return self._parse_markdown(file_path)
            elif file_type == "text":
                return self._parse_text(file_path)
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
        except FileParseError:
            raise
        except Exception as e:
            logger.error(f"Failed to parse file {file_path}: {e}")
            raise FileParseError(f"Failed to parse file: {e}")

    def parse_file_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Parse file from bytes content.

        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to determine file type)

        Returns:
            Extracted text content

        Raises:
            UnsupportedFileTypeError: If file type is not supported
            FileParseError: If parsing fails
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {ext}. Supported types: {list(self.SUPPORTED_EXTENSIONS.keys())}"
            )

        if ext not in self.allowed_extensions:
            raise UnsupportedFileTypeError(
                f"File type {ext} is not in the allowed list: {self.allowed_extensions}"
            )

        file_type = self.SUPPORTED_EXTENSIONS[ext]

        try:
            if file_type == "word":
                return self._parse_word_bytes(file_bytes)
            elif file_type == "pdf":
                return self._parse_pdf_bytes(file_bytes)
            elif file_type == "markdown":
                return file_bytes.decode("utf-8", errors="ignore")
            elif file_type == "text":
                return file_bytes.decode("utf-8", errors="ignore")
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
        except FileParseError:
            raise
        except Exception as e:
            logger.error(f"Failed to parse file bytes for {filename}: {e}")
            raise FileParseError(f"Failed to parse file: {e}")

    def _parse_word(self, file_path: str) -> str:
        """Parse Word document (.docx)."""
        try:
            doc = WordDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            raise FileParseError(f"Failed to parse Word document: {e}")

    def _parse_word_bytes(self, file_bytes: bytes) -> str:
        """Parse Word document from bytes."""
        import io

        try:
            doc = WordDocument(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            raise FileParseError(f"Failed to parse Word document from bytes: {e}")

    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF document (.pdf)."""
        try:
            doc = PDFDocument(file_path)
            text_parts = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text_parts.append(page.get_text())

            return "\n\n".join(text_parts)
        except Exception as e:
            raise FileParseError(f"Failed to parse PDF document: {e}")

    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """Parse PDF document from bytes."""
        try:
            doc = PDFDocument(stream=file_bytes)
            text_parts = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text_parts.append(page.get_text())

            return "\n\n".join(text_parts)
        except Exception as e:
            raise FileParseError(f"Failed to parse PDF document from bytes: {e}")

    def _parse_markdown(self, file_path: str) -> str:
        """Parse Markdown file (.md)."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    return f.read()
            except Exception as e:
                raise FileParseError(f"Failed to read Markdown file: {e}")
        except Exception as e:
            raise FileParseError(f"Failed to parse Markdown file: {e}")

    def _parse_text(self, file_path: str) -> str:
        """Parse plain text file (.txt)."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    return f.read()
            except Exception as e:
                raise FileParseError(f"Failed to read text file: {e}")
        except Exception as e:
            raise FileParseError(f"Failed to parse text file: {e}")

    def is_supported(self, filename: str) -> bool:
        """
        Check if a file type is supported.

        Args:
            filename: Name of the file to check

        Returns:
            True if supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS and ext in self.allowed_extensions

    @classmethod
    def get_file_type(cls, filename: str) -> Optional[str]:
        """
        Get the file type for a filename.

        Args:
            filename: Name of the file

        Returns:
            File type string or None if not supported
        """
        ext = Path(filename).suffix.lower()
        return cls.SUPPORTED_EXTENSIONS.get(ext)


# ============================================================================
# Convenience functions
# ============================================================================


def parse_file(file_path: str, allowed_extensions: Optional[List[str]] = None) -> str:
    """
    Convenience function to parse a file.

    Args:
        file_path: Path to the file to parse
        allowed_extensions: List of allowed file extensions

    Returns:
        Extracted text content
    """
    parser = FileParser(allowed_extensions)
    return parser.parse_file(file_path)


def parse_file_bytes(
    file_bytes: bytes, filename: str, allowed_extensions: Optional[List[str]] = None
) -> str:
    """
    Convenience function to parse file from bytes.

    Args:
        file_bytes: File content as bytes
        filename: Original filename
        allowed_extensions: List of allowed file extensions

    Returns:
        Extracted text content
    """
    parser = FileParser(allowed_extensions)
    return parser.parse_file_bytes(file_bytes, filename)


# ============================================================================
# Export
# ============================================================================

__all__ = [
    "FileParser",
    "FileParseError",
    "UnsupportedFileTypeError",
    "FileNotFoundError",
    "parse_file",
    "parse_file_bytes",
]
